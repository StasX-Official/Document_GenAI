#!/usr/bin/env python3
import argparse
import asyncio
import json
import logging
import os
import sys
import time
import uuid
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

try:
    from rich.console import Console
    from rich.logging import RichHandler
    from rich.progress import Progress, SpinnerColumn, TextColumn
    from rich.prompt import Prompt
    from rich import print as rprint
    RICH_AVAILABLE = True
except ImportError:
    RICH_AVAILABLE = False

import google.generativeai as genai
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches


class ModelType(str, Enum):
    GEMINI_PRO = "gemini-pro"
    GEMINI_FLASH = "gemini-2.0-flash"
    GEMINI_PRO_VISION = "gemini-pro-vision"


class Config:
    DEFAULT_CONFIG = {
        "api_key": "",
        "default_model": ModelType.GEMINI_FLASH.value,
        "output_directory": "generated_documents",
        "document_template": "default",
        "log_level": "INFO",
        "history_enabled": True,
        "max_history_items": 50,
        "default_temperature": 0.7,
    }
    
    def __init__(self):
        self.config_dir = Path.home() / ".gemini_docs"
        self.config_path = self.config_dir / "config.json"
        self.history_path = self.config_dir / "history.json"
        self.settings = self._load_config()
        
    def _load_config(self) -> Dict[str, Any]:
        self.config_dir.mkdir(exist_ok=True)
        
        if not self.config_path.exists():
            config = self.DEFAULT_CONFIG.copy()
            config["api_key"] = os.environ.get("GEMINI_API_KEY", "")
            
            with open(self.config_path, "w") as f:
                json.dump(config, f, indent=2)
            
            return config
        
        try:
            with open(self.config_path, "r") as f:
                config = json.load(f)
                
            for key, value in self.DEFAULT_CONFIG.items():
                if key not in config:
                    config[key] = value
                    
            env_api_key = os.environ.get("GEMINI_API_KEY")
            if env_api_key:
                config["api_key"] = env_api_key
                
            return config
        except Exception as e:
            logging.error(f"Failed to load config: {e}. Using defaults.")
            return self.DEFAULT_CONFIG.copy()
            
    def save(self):
        with open(self.config_path, "w") as f:
            json.dump(self.settings, f, indent=2)
            
    def get_output_dir(self) -> Path:
        output_path = Path(self.settings["output_directory"])
        if not output_path.is_absolute():
            output_path = Path.cwd() / output_path
        output_path.mkdir(exist_ok=True, parents=True)
        return output_path
    
    def add_to_history(self, prompt: str, filepath: str) -> None:
        if not self.settings.get("history_enabled", True):
            return
            
        try:
            history = []
            if self.history_path.exists():
                with open(self.history_path, "r") as f:
                    history = json.load(f)
                    
            history.append({
                "timestamp": datetime.now().isoformat(),
                "prompt": prompt,
                "filepath": filepath,
            })
            
            max_items = self.settings.get("max_history_items", 50)
            history = history[-max_items:] if len(history) > max_items else history
            
            with open(self.history_path, "w") as f:
                json.dump(history, f, indent=2)
        except Exception as e:
            logging.warning(f"Failed to update history: {e}")


class DocumentTemplate:
    def apply_template(self, doc: Document, content: str, prompt: str) -> Document:
        for i in range(len(doc.paragraphs)-1, -1, -1):
            p = doc.paragraphs[i]
            p._element.getparent().remove(p._element)
            
        header = doc.add_heading("AI Generated Document", level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading("About this document", level=2)
        metadata = doc.add_paragraph()
        metadata.add_run("Generated on: ").bold = True
        metadata.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        metadata.add_run("\nPrompt: ").bold = True
        prompt_run = metadata.add_run(prompt)
        prompt_run.italic = True
        
        doc.add_paragraph("").paragraph_format.space_after = Pt(12)
        
        doc.add_heading("Generated Content", level=2)
        
        for paragraph_text in content.split('\n\n'):
            if not paragraph_text.strip():
                continue
                
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.first_line_indent = Inches(0.25)
            p.add_run(paragraph_text)
            
        doc.add_paragraph("").paragraph_format.space_before = Pt(20)
        footer = doc.add_paragraph("Generated using Gemini AI")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.runs[0]
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(100, 100, 100)
        
        return doc


class DocumentGenerator:
    def __init__(self):
        self._setup_logging()
        self.config = Config()
        self.template = DocumentTemplate()
        self.console = Console() if RICH_AVAILABLE else None
        self._init_gemini_client()
    
    def _setup_logging(self):
        log_level = os.environ.get("LOG_LEVEL", "INFO")
        
        if RICH_AVAILABLE:
            logging.basicConfig(
                level=log_level,
                format="%(message)s",
                datefmt="[%X]",
                handlers=[RichHandler(rich_tracebacks=True)]
            )
        else:
            logging.basicConfig(
                level=log_level,
                format="%(asctime)s - %(name)s - %(levelname)s - %(message)s"
            )
        
        self.logger = logging.getLogger("gemini-docs")
        
    def _init_gemini_client(self):
        api_key = self.config.settings["api_key"]
        
        if not api_key:
            self.logger.warning("API key not set. You will be prompted to enter it.")
            return
            
        try:
            genai.configure(api_key=api_key)
            self.logger.debug("Gemini API client initialized successfully")
        except Exception as e:
            self.logger.error(f"Failed to initialize Gemini client: {e}")
            
    def create_document(self, content: str, prompt: str) -> Tuple[str, str]:
        timestamp = int(time.time())
        filename = f"document_{timestamp}_{uuid.uuid4().hex[:6]}.docx"
        
        output_dir = self.config.get_output_dir()
        filepath = output_dir / filename
        
        doc = Document()
        doc = self.template.apply_template(doc, content, prompt)
        doc.save(filepath)
        
        try:
            rel_path = filepath.relative_to(Path.cwd())
            display_path = str(rel_path)
        except ValueError:
            display_path = str(filepath)
            
        return str(filepath), display_path
    
    async def generate_content(self, prompt: str) -> Optional[str]:
        api_key = self.config.settings["api_key"]
        
        if not api_key:
            if RICH_AVAILABLE:
                api_key = Prompt.ask("Enter your Gemini API key")
            else:
                api_key = input("Enter your Gemini API key: ")
                
            if not api_key:
                self.logger.error("API key is required to continue")
                return None
                
            self.config.settings["api_key"] = api_key
            self.config.save()
            genai.configure(api_key=api_key)
        
        model_name = self.config.settings.get("default_model", ModelType.GEMINI_FLASH.value)
        temperature = float(self.config.settings.get("default_temperature", 0.7))
        
        try:
            if RICH_AVAILABLE and self.console:
                with Progress(
                    SpinnerColumn(),
                    TextColumn("[bold blue]Generating content..."),
                    console=self.console
                ) as progress:
                    task = progress.add_task("Generating", total=None)
                    
                    model = genai.GenerativeModel(
                        model_name=model_name,
                        generation_config={"temperature": temperature}
                    )
                    response = await asyncio.to_thread(
                        model.generate_content, prompt
                    )
                    progress.update(task, completed=True)
            else:
                print("Generating content...")
                model = genai.GenerativeModel(
                    model_name=model_name,
                    generation_config={"temperature": temperature}
                )
                response = model.generate_content(prompt)
            
            if not response or not hasattr(response, "text"):
                raise RuntimeError("Failed to generate content: Empty response")
                
            return response.text
                
        except Exception as e:
            self.logger.error(f"Error generating content: {e}")
            return None

    async def process_prompt(self, prompt: str) -> bool:
        try:
            content = await self.generate_content(prompt)
            if not content:
                return False
                
            filepath, display_path = self.create_document(content, prompt)
            
            self.config.add_to_history(prompt, filepath)
            
            if RICH_AVAILABLE and self.console:
                self.console.print(f"[bold green]✓[/] Document saved: [link=file://{filepath}]{display_path}[/link]")
            else:
                print(f"✓ Document saved: {display_path}")
                
            return True
            
        except Exception as e:
            self.logger.error(f"Failed to process prompt: {e}")
            return False
    
    async def run_cli(self) -> None:
        if RICH_AVAILABLE:
            self.console.print("[bold]Gemini Document Generator[/bold] [dim]v1.1.0[/dim]")
            self.console.print("Type [bold cyan]exit[/] to quit or enter a prompt to generate a document.")
        else:
            print("Gemini Document Generator v1.1.0")
            print("Type 'exit' to quit or enter a prompt to generate a document.")
        
        while True:
            try:
                if RICH_AVAILABLE:
                    user_input = Prompt.ask("\nPrompt")
                else:
                    user_input = input("\nPrompt: ")
                    
                if user_input.lower() in ("exit", "quit", "q"):
                    if RICH_AVAILABLE:
                        self.console.print("[dim]Exiting...[/dim]")
                    else:
                        print("Exiting...")
                    break
                    
                await self.process_prompt(user_input)
                
            except KeyboardInterrupt:
                print("\nOperation cancelled by user")
                break
            except Exception as e:
                self.logger.error(f"Error: {e}")


async def async_main(args: argparse.Namespace) -> None:
    generator = DocumentGenerator()
    
    if args.prompt:
        await generator.process_prompt(args.prompt)
        return
        
    await generator.run_cli()


def main():
    parser = argparse.ArgumentParser(description="Gemini Document Generator")
    parser.add_argument("-p", "--prompt", help="Generate a document with this prompt and exit")
    parser.add_argument("--config", help="Path to custom config file")
    parser.add_argument("--debug", action="store_true", help="Enable debug logging")
    args = parser.parse_args()
    
    if args.debug:
        logging.getLogger().setLevel(logging.DEBUG)
    
    try:
        if sys.platform == "win32":
            asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
            
        asyncio.run(async_main(args))
    except Exception as e:
        logging.error(f"Application error: {e}", exc_info=True)
    finally:
        print("Shutting down...")
        time.sleep(0.5)


if __name__ == "__main__":
    main()
